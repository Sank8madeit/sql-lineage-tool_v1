import streamlit as st
import pandas as pd
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sqlglot
from sqlglot import exp
import io
from datetime import datetime
import re

# ---------------------------
# PAGE CONFIG
# ---------------------------
st.set_page_config(
    page_title="SQL Docs Pro",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ---------------------------
# CUSTOM CSS
# ---------------------------
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=IBM+Plex+Sans:wght@300;400;500;600&display=swap');

    html, body, [class*="css"] {
        font-family: 'IBM Plex Sans', sans-serif;
    }

    .stApp {
        background: #0f1117;
        color: #e2e8f0;
    }

    /* Header */
    .hero-header {
        background: linear-gradient(135deg, #0f1117 0%, #1a1f2e 50%, #0f1117 100%);
        border-bottom: 1px solid #2d3748;
        padding: 2rem 0 1.5rem 0;
        margin-bottom: 2rem;
    }
    .hero-title {
        font-family: 'IBM Plex Mono', monospace;
        font-size: 1.8rem;
        font-weight: 600;
        color: #63b3ed;
        letter-spacing: -0.5px;
        margin: 0;
    }
    .hero-subtitle {
        font-size: 0.85rem;
        color: #718096;
        margin-top: 0.25rem;
        font-weight: 300;
    }
    .version-badge {
        display: inline-block;
        background: #1a365d;
        color: #63b3ed;
        font-family: 'IBM Plex Mono', monospace;
        font-size: 0.7rem;
        padding: 2px 8px;
        border-radius: 3px;
        border: 1px solid #2b6cb0;
        margin-left: 10px;
        vertical-align: middle;
    }

    /* Sidebar-style config panel */
    .config-panel {
        background: #1a1f2e;
        border: 1px solid #2d3748;
        border-radius: 8px;
        padding: 1.25rem;
        margin-bottom: 1rem;
    }
    .config-label {
        font-size: 0.7rem;
        font-weight: 600;
        color: #718096;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 0.5rem;
    }

    /* SQL Editor */
    .stTextArea textarea {
        font-family: 'IBM Plex Mono', monospace !important;
        font-size: 0.82rem !important;
        background: #0d1117 !important;
        color: #a8d8a8 !important;
        border: 1px solid #2d3748 !important;
        border-radius: 6px !important;
        line-height: 1.6 !important;
    }
    .stTextArea textarea:focus {
        border-color: #4299e1 !important;
        box-shadow: 0 0 0 2px rgba(66, 153, 225, 0.15) !important;
    }

    /* Buttons */
    .stButton > button {
        background: #2b6cb0 !important;
        color: #ebf8ff !important;
        border: none !important;
        border-radius: 6px !important;
        font-family: 'IBM Plex Sans', sans-serif !important;
        font-weight: 500 !important;
        font-size: 0.85rem !important;
        padding: 0.5rem 1.5rem !important;
        transition: all 0.2s !important;
        letter-spacing: 0.3px;
    }
    .stButton > button:hover {
        background: #3182ce !important;
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(66, 153, 225, 0.3) !important;
    }

    /* Download buttons */
    .stDownloadButton > button {
        background: #1a202c !important;
        color: #68d391 !important;
        border: 1px solid #2d3748 !important;
        border-radius: 6px !important;
        font-family: 'IBM Plex Mono', monospace !important;
        font-size: 0.78rem !important;
        padding: 0.4rem 1rem !important;
        transition: all 0.2s !important;
    }
    .stDownloadButton > button:hover {
        background: #2d3748 !important;
        border-color: #68d391 !important;
    }

    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        background: #1a1f2e !important;
        border-radius: 8px 8px 0 0 !important;
        border-bottom: 1px solid #2d3748 !important;
        gap: 0 !important;
    }
    .stTabs [data-baseweb="tab"] {
        font-family: 'IBM Plex Sans', sans-serif !important;
        font-size: 0.8rem !important;
        font-weight: 500 !important;
        color: #718096 !important;
        padding: 0.6rem 1.2rem !important;
        border-radius: 0 !important;
    }
    .stTabs [aria-selected="true"] {
        background: #2d3748 !important;
        color: #63b3ed !important;
        border-bottom: 2px solid #4299e1 !important;
    }

    /* Metrics */
    [data-testid="metric-container"] {
        background: #1a1f2e !important;
        border: 1px solid #2d3748 !important;
        border-radius: 8px !important;
        padding: 0.75rem !important;
    }
    [data-testid="metric-container"] label {
        color: #718096 !important;
        font-size: 0.72rem !important;
        text-transform: uppercase;
        letter-spacing: 0.8px;
    }
    [data-testid="metric-container"] [data-testid="stMetricValue"] {
        color: #63b3ed !important;
        font-family: 'IBM Plex Mono', monospace !important;
        font-size: 1.4rem !important;
    }

    /* Dataframe */
    .stDataFrame {
        border: 1px solid #2d3748 !important;
        border-radius: 6px !important;
    }

    /* Status messages */
    .stSuccess {
        background: #1a2e1a !important;
        border: 1px solid #276749 !important;
        border-radius: 6px !important;
    }
    .stError {
        background: #2d1515 !important;
        border: 1px solid #742a2a !important;
        border-radius: 6px !important;
    }

    /* Section headers */
    .section-header {
        font-family: 'IBM Plex Mono', monospace;
        font-size: 0.75rem;
        font-weight: 600;
        color: #4299e1;
        text-transform: uppercase;
        letter-spacing: 1.5px;
        border-bottom: 1px solid #2d3748;
        padding-bottom: 0.5rem;
        margin-bottom: 1rem;
    }

    /* Doc cards */
    .doc-card {
        background: #1a1f2e;
        border: 1px solid #2d3748;
        border-radius: 8px;
        padding: 1rem;
        margin-bottom: 0.75rem;
        border-left: 3px solid #4299e1;
    }
    .doc-card-title {
        font-size: 0.82rem;
        font-weight: 600;
        color: #e2e8f0;
        margin-bottom: 0.25rem;
    }
    .doc-card-desc {
        font-size: 0.75rem;
        color: #718096;
    }

    /* Checkbox */
    .stCheckbox label {
        color: #a0aec0 !important;
        font-size: 0.82rem !important;
    }

    /* Selectbox */
    .stSelectbox label {
        color: #718096 !important;
        font-size: 0.75rem !important;
        text-transform: uppercase;
        letter-spacing: 0.8px;
    }

    /* Text input */
    .stTextInput label {
        color: #718096 !important;
        font-size: 0.75rem !important;
        text-transform: uppercase;
        letter-spacing: 0.8px;
    }
    .stTextInput input {
        background: #0d1117 !important;
        border: 1px solid #2d3748 !important;
        color: #e2e8f0 !important;
        border-radius: 6px !important;
        font-size: 0.85rem !important;
    }

    /* Divider */
    hr {
        border-color: #2d3748 !important;
    }

    /* Expander */
    .streamlit-expanderHeader {
        background: #1a1f2e !important;
        border-radius: 6px !important;
        font-size: 0.82rem !important;
        color: #a0aec0 !important;
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------
# HEADER
# ---------------------------
st.markdown("""
<div class="hero-header">
    <div class="hero-title">
        ⬡ SQL DOCS PRO
        <span class="version-badge">v5.0</span>
    </div>
    <div class="hero-subtitle">Source-to-Target Mapping · Data Lineage · ODD · FRD · Data Dictionary</div>
</div>
""", unsafe_allow_html=True)

# ---------------------------
# LINEAGE ENGINE
# ---------------------------

class LineageEngine:
    def __init__(self, sql):
        self.sql = sql
        self.tree = None
        self.alias_map = {}
        self.cte_map = {}
        self.lineage = []

    def parse_sql(self):
        try:
            self.tree = sqlglot.parse_one(self.sql, read="oracle")
        except Exception as e:
            raise Exception(f"SQL Parsing Failed: {e}")

    def build_alias_map(self):
        for table in self.tree.find_all(exp.Table):
            name = table.name
            alias = table.alias
            if alias:
                self.alias_map[alias] = name
            else:
                self.alias_map[name] = name

    def extract_ctes(self):
        for cte in self.tree.find_all(exp.CTE):
            self.cte_map[cte.alias] = cte.this

    def resolve_table(self, table_name):
        if table_name in self.alias_map:
            return self.alias_map[table_name]
        if table_name in self.cte_map:
            cte_query = self.cte_map[table_name]
            for t in cte_query.find_all(exp.Table):
                return t.name
        return table_name

    def detect_transformation_type(self, expr_sql):
        expr_upper = expr_sql.upper().strip()
        if re.search(r'\bCASE\b', expr_upper):
            return "Conditional"
        elif re.search(r'\bCONCAT\b|\|\|', expr_upper):
            return "Concatenation"
        elif re.search(r'\bCOALESCE\b|\bNVL\b|\bNULLIF\b', expr_upper):
            return "Null Handling"
        elif re.search(r'\bTO_DATE\b|\bTO_CHAR\b|\bTRUNC\b|\bEXTRACT\b|\bDATEADD\b|\bDATEDIFF\b', expr_upper):
            return "Date/Type Cast"
        elif re.search(r'\bSUM\b|\bAVG\b|\bCOUNT\b|\bMAX\b|\bMIN\b', expr_upper):
            return "Aggregation"
        elif re.search(r'\bUPPER\b|\bLOWER\b|\bTRIM\b|\bSUBSTR\b|\bREPLACE\b', expr_upper):
            return "String Function"
        elif re.search(r'[+\-\*\/]', expr_upper):
            return "Arithmetic"
        elif re.match(r"^'.*'$", expr_upper) or re.match(r"^\d+$", expr_upper):
            return "Literal/Constant"
        elif expr_upper == expr_sql.upper() and '(' not in expr_sql:
            return "Direct Mapping"
        else:
            return "Expression"

    def process_select(self, select, target_table="FINAL_OUTPUT", depth=0):
        for proj in select.expressions:
            target = proj.alias_or_name
            expr = proj.this
            if not expr:
                continue

            expr_sql = proj.sql()
            transform_type = self.detect_transformation_type(expr_sql)
            cols = list(expr.find_all(exp.Column))

            if not cols:
                self.lineage.append({
                    "SOURCE_COLUMN": "N/A",
                    "SOURCE_TABLE": "N/A",
                    "TRANSFORMATION_LOGIC": expr.sql(),
                    "TRANSFORMATION_TYPE": transform_type,
                    "TARGET_COLUMN": target,
                    "TARGET_TABLE": target_table
                })
            else:
                for col in cols:
                    table_alias = col.table
                    real_table = self.resolve_table(table_alias) if table_alias else "UNKNOWN"
                    self.lineage.append({
                        "SOURCE_COLUMN": col.name,
                        "SOURCE_TABLE": real_table,
                        "TRANSFORMATION_LOGIC": expr.sql(),
                        "TRANSFORMATION_TYPE": transform_type,
                        "TARGET_COLUMN": target,
                        "TARGET_TABLE": target_table
                    })

    def run(self):
        self.parse_sql()
        self.build_alias_map()
        self.extract_ctes()
        for node in self.tree.walk():
            if isinstance(node, exp.Select):
                self.process_select(node)
        return self.lineage

    def get_source_tables(self):
        tables = set()
        for item in self.lineage:
            if item["SOURCE_TABLE"] not in ("N/A", "UNKNOWN", ""):
                tables.add(item["SOURCE_TABLE"])
        return sorted(tables)

    def get_target_columns(self):
        return sorted(set(i["TARGET_COLUMN"] for i in self.lineage if i["TARGET_COLUMN"]))


# ---------------------------
# DOC HELPERS
# ---------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def style_table_header(table, header_color="1A365D", text_color="FFFFFF"):
    for cell in table.rows[0].cells:
        set_cell_bg(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor.from_string(text_color)
                run.font.bold = True
                run.font.size = Pt(9)

def add_table_from_df(doc, df, header_color="1A365D"):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = col
        set_cell_bg(hdr[i], header_color)
        for para in hdr[i].paragraphs:
            run = para.runs[0] if para.runs else para.add_run(col)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8.5)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    row_colors = ["FFFFFF", "F7FAFC"]
    for row_idx, row_data in df.iterrows():
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val) if val is not None else ""
            set_cell_bg(cell, row_colors[row_idx % 2])
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
    return table

def add_section_title(doc, title, level=1):
    colors = {1: "1A365D", 2: "2C5282", 3: "2B6CB0"}
    p = doc.add_heading(title, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(colors.get(level, "1A365D"))
    return p

def add_info_box(doc, label, value):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    c1, c2 = tbl.rows[0].cells
    set_cell_bg(c1, "EBF8FF")
    set_cell_bg(c2, "FFFFFF")
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(label)
    r1.bold = True
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor(26, 54, 93)
    c2.paragraphs[0].add_run(str(value)).font.size = Pt(8.5)
    return tbl

def doc_to_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def base_doc(title_text, subtitle="", project="", author=""):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title_text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 54, 93)

    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.runs[0].font.size = Pt(11)
        p2.runs[0].font.color.rgb = RGBColor(113, 128, 150)

    # Meta line
    meta = []
    if project:
        meta.append(f"Project: {project}")
    meta.append(f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')}")
    if author:
        meta.append(f"Author: {author}")

    if meta:
        pm = doc.add_paragraph(" · ".join(meta))
        pm.runs[0].font.size = Pt(8.5)
        pm.runs[0].font.color.rgb = RGBColor(160, 174, 192)

    doc.add_paragraph()
    return doc


# ---------------------------
# DOCUMENT GENERATORS
# ---------------------------

def generate_sttm_doc(df, project, author):
    doc = base_doc("Source-to-Target Mapping (STTM)", "Column-level data lineage mapping", project, author)
    add_section_title(doc, "STTM Matrix", 1)
    doc.add_paragraph("The table below maps every source column and transformation to its corresponding target column.")
    doc.add_paragraph()
    add_table_from_df(doc, df[["SOURCE_TABLE", "SOURCE_COLUMN", "TRANSFORMATION_TYPE", "TRANSFORMATION_LOGIC", "TARGET_COLUMN", "TARGET_TABLE"]])
    return doc_to_bytes(doc)


def generate_data_dict_doc(df, project, author):
    doc = base_doc("Data Dictionary", "Column definitions and transformation metadata", project, author)
    add_section_title(doc, "Column Reference", 1)

    seen = set()
    for _, row in df.iterrows():
        col = row["TARGET_COLUMN"]
        if col in seen:
            continue
        seen.add(col)

        p = doc.add_paragraph()
        r = p.add_run(f"  {col}")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(26, 86, 160)

        rows_for_col = df[df["TARGET_COLUMN"] == col]
        sources = rows_for_col["SOURCE_TABLE"].unique()
        src_cols = rows_for_col["SOURCE_COLUMN"].unique()
        trans_type = rows_for_col["TRANSFORMATION_TYPE"].iloc[0]
        logic = rows_for_col["TRANSFORMATION_LOGIC"].iloc[0]

        info = [
            ("Source Table(s)", ", ".join(s for s in sources if s != "N/A") or "Derived"),
            ("Source Column(s)", ", ".join(s for s in src_cols if s != "N/A") or "N/A"),
            ("Transformation Type", trans_type),
            ("Transformation Logic", logic),
            ("Data Type", "VARCHAR2 / NUMBER (inferred)"),
            ("Nullable", "Yes"),
            ("Description", f"Target column '{col}' derived via {trans_type.lower()} logic."),
        ]
        tbl = doc.add_table(rows=len(info), cols=2)
        tbl.style = 'Table Grid'
        for i, (k, v) in enumerate(info):
            c1, c2 = tbl.rows[i].cells
            set_cell_bg(c1, "EBF8FF" if i % 2 == 0 else "F0F4F8")
            set_cell_bg(c2, "FFFFFF" if i % 2 == 0 else "FAFAFA")
            r1 = c1.paragraphs[0].add_run(k)
            r1.bold = True
            r1.font.size = Pt(8)
            r1.font.color.rgb = RGBColor(26, 54, 93)
            c2.paragraphs[0].add_run(v).font.size = Pt(8)

        doc.add_paragraph()
    return doc_to_bytes(doc)


def generate_lineage_doc(df, engine, project, author):
    doc = base_doc("Data Lineage Document", "End-to-end data flow and transformation chain", project, author)

    add_section_title(doc, "1. Overview", 1)
    source_tables = engine.get_source_tables()
    target_cols = engine.get_target_columns()

    doc.add_paragraph(
        f"This document describes the data lineage for the SQL transformation. "
        f"Data originates from {len(source_tables)} source table(s) and produces "
        f"{len(target_cols)} target column(s) in the FINAL_OUTPUT."
    )

    add_section_title(doc, "2. Source Tables", 2)
    for t in source_tables:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(t).font.bold = True

    add_section_title(doc, "3. Transformation Summary", 2)
    type_counts = df["TRANSFORMATION_TYPE"].value_counts().reset_index()
    type_counts.columns = ["Transformation Type", "Count"]
    add_table_from_df(doc, type_counts, "2C5282")
    doc.add_paragraph()

    add_section_title(doc, "4. Column-Level Lineage", 2)
    add_table_from_df(doc, df[["SOURCE_TABLE", "SOURCE_COLUMN", "TRANSFORMATION_TYPE", "TARGET_COLUMN"]])
    doc.add_paragraph()

    add_section_title(doc, "5. CTE / Subquery Chain", 2)
    cte_names = list(engine.cte_map.keys())
    if cte_names:
        doc.add_paragraph("The following CTEs / subqueries were identified:")
        for c in cte_names:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(c).font.bold = True
    else:
        doc.add_paragraph("No CTEs or subqueries detected.")

    return doc_to_bytes(doc)


def generate_odd_doc(df, engine, project, author):
    doc = base_doc("Output Data Definition (ODD)", "Specification of output dataset structure and rules", project, author)

    add_section_title(doc, "1. Purpose", 1)
    doc.add_paragraph(
        "This Output Data Definition document describes the structure, data types, "
        "business rules, and constraints of the output dataset produced by the SQL transformation."
    )

    add_section_title(doc, "2. Output Dataset Summary", 2)
    summary_data = {
        "Attribute": ["Dataset Name", "Source System(s)", "Number of Output Columns", "Transformation Count", "Generated On"],
        "Value": [
            "FINAL_OUTPUT",
            ", ".join(engine.get_source_tables()) or "Unknown",
            str(len(engine.get_target_columns())),
            str(len(df)),
            datetime.now().strftime("%d %b %Y")
        ]
    }
    add_table_from_df(doc, pd.DataFrame(summary_data), "1A365D")
    doc.add_paragraph()

    add_section_title(doc, "3. Column Specifications", 2)
    seen = set()
    specs = []
    for _, row in df.iterrows():
        col = row["TARGET_COLUMN"]
        if col in seen:
            continue
        seen.add(col)
        src_tables = df[df["TARGET_COLUMN"] == col]["SOURCE_TABLE"].unique()
        specs.append({
            "Column Name": col,
            "Data Type": "VARCHAR2 / NUMBER",
            "Nullable": "Y",
            "Source Table(s)": ", ".join(s for s in src_tables if s not in ("N/A", "UNKNOWN")) or "Derived",
            "Transform Type": row["TRANSFORMATION_TYPE"],
            "Business Rule": f"Derived via {row['TRANSFORMATION_TYPE'].lower()}",
            "PK/FK": "N/A"
        })
    add_table_from_df(doc, pd.DataFrame(specs), "1A365D")
    doc.add_paragraph()

    add_section_title(doc, "4. Data Quality Rules", 2)
    rules = [
        ("DQ-001", "Completeness", "All mandatory fields must be non-null", "Critical"),
        ("DQ-002", "Uniqueness", "Target column values must align with source key uniqueness", "High"),
        ("DQ-003", "Referential Integrity", "Foreign key references must resolve to valid source records", "High"),
        ("DQ-004", "Format Validity", "Date/numeric types must conform to expected formats", "Medium"),
        ("DQ-005", "Transformation Accuracy", "Transformation logic output must match expected business values", "Critical"),
    ]
    add_table_from_df(doc, pd.DataFrame(rules, columns=["Rule ID", "Category", "Description", "Severity"]), "2C5282")

    return doc_to_bytes(doc)


def generate_frd_doc(df, engine, project, author):
    doc = base_doc("Functional Requirements Document (FRD)", "Business and functional specification for SQL transformation", project, author)

    add_section_title(doc, "1. Introduction", 1)
    doc.add_paragraph(
        "This Functional Requirements Document captures the business requirements, "
        "functional specifications, and acceptance criteria for the SQL transformation pipeline."
    )

    add_section_title(doc, "1.1 Scope", 2)
    doc.add_paragraph(
        f"The scope of this document covers the data transformation from "
        f"{len(engine.get_source_tables())} source table(s) to the FINAL_OUTPUT dataset "
        f"containing {len(engine.get_target_columns())} columns."
    )

    add_section_title(doc, "1.2 Stakeholders", 2)
    stakeholders = pd.DataFrame({
        "Role": ["Data Engineer", "Business Analyst", "Data Architect", "QA Engineer"],
        "Responsibility": ["Implement transformation", "Define business rules", "Review lineage", "Validate output"],
        "Sign-off Required": ["Yes", "Yes", "Yes", "No"]
    })
    add_table_from_df(doc, stakeholders, "1A365D")
    doc.add_paragraph()

    add_section_title(doc, "2. Source Systems", 1)
    for t in engine.get_source_tables():
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(t)
        run.bold = True

    add_section_title(doc, "3. Functional Requirements", 1)
    reqs = []
    for i, col in enumerate(engine.get_target_columns(), 1):
        rows_for_col = df[df["TARGET_COLUMN"] == col]
        trans_type = rows_for_col["TRANSFORMATION_TYPE"].iloc[0]
        src_tables = rows_for_col["SOURCE_TABLE"].unique()
        reqs.append({
            "FR-ID": f"FR-{i:03d}",
            "Target Column": col,
            "Requirement": f"Populate '{col}' using {trans_type.lower()} from {', '.join(s for s in src_tables if s != 'N/A') or 'expression'}",
            "Priority": "High" if trans_type == "Direct Mapping" else "Medium",
            "Status": "Defined"
        })
    add_table_from_df(doc, pd.DataFrame(reqs), "1A365D")
    doc.add_paragraph()

    add_section_title(doc, "4. Non-Functional Requirements", 1)
    nfrs = pd.DataFrame({
        "NFR-ID": ["NFR-001", "NFR-002", "NFR-003", "NFR-004"],
        "Category": ["Performance", "Scalability", "Data Quality", "Auditability"],
        "Requirement": [
            "Transformation must complete within SLA window",
            "Must handle incremental and full-load patterns",
            "Zero tolerance for data loss on critical fields",
            "All transformations must be logged for audit trail"
        ],
        "Priority": ["High", "Medium", "Critical", "High"]
    })
    add_table_from_df(doc, nfrs, "2C5282")
    doc.add_paragraph()

    add_section_title(doc, "5. Acceptance Criteria", 1)
    criteria = pd.DataFrame({
        "AC-ID": ["AC-001", "AC-002", "AC-003", "AC-004", "AC-005"],
        "Criterion": [
            "All target columns defined in FRD are present in output",
            "Row count in output matches expected source record count",
            "Transformation logic matches documented business rules",
            "Data dictionary covers all output columns",
            "Lineage traceable from source to target for every column"
        ],
        "Test Method": ["Schema check", "Count assertion", "Spot check", "Doc review", "Lineage trace"],
        "Pass/Fail": ["—"] * 5
    })
    add_table_from_df(doc, criteria, "276749")

    return doc_to_bytes(doc)


def generate_excel(df):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as writer:
        # STTM
        sttm_df = df[["SOURCE_TABLE", "SOURCE_COLUMN", "TRANSFORMATION_TYPE", "TRANSFORMATION_LOGIC", "TARGET_COLUMN", "TARGET_TABLE"]]
        sttm_df.to_excel(writer, sheet_name="STTM", index=False)

        # Data Dictionary
        seen = set()
        dict_rows = []
        for _, row in df.iterrows():
            col = row["TARGET_COLUMN"]
            if col in seen:
                continue
            seen.add(col)
            src_tables = df[df["TARGET_COLUMN"] == col]["SOURCE_TABLE"].unique()
            dict_rows.append({
                "Column Name": col,
                "Source Table(s)": ", ".join(s for s in src_tables if s not in ("N/A", "UNKNOWN")) or "Derived",
                "Transformation Type": row["TRANSFORMATION_TYPE"],
                "Transformation Logic": row["TRANSFORMATION_LOGIC"],
                "Data Type": "VARCHAR2 / NUMBER",
                "Nullable": "Y",
                "Description": f"Derived via {row['TRANSFORMATION_TYPE'].lower()}"
            })
        pd.DataFrame(dict_rows).to_excel(writer, sheet_name="DataDictionary", index=False)

        # Lineage summary
        type_counts = df["TRANSFORMATION_TYPE"].value_counts().reset_index()
        type_counts.columns = ["Type", "Count"]
        type_counts.to_excel(writer, sheet_name="LineageSummary", index=False)

    buf.seek(0)
    return buf.read()


# ---------------------------
# LAYOUT
# ---------------------------

col_left, col_right = st.columns([3, 2], gap="large")

with col_left:
    st.markdown('<div class="section-header">SQL Input</div>', unsafe_allow_html=True)
    sql_input = st.text_area(
        label="sql_editor",
        placeholder="-- Paste your Oracle / ANSI SQL here\nSELECT\n    a.customer_id,\n    UPPER(b.name) AS customer_name\nFROM customers a\nJOIN orders b ON a.id = b.customer_id",
        height=320,
        label_visibility="collapsed"
    )

with col_right:
    st.markdown('<div class="section-header">Configuration</div>', unsafe_allow_html=True)

    project_name = st.text_input("Project / System Name", placeholder="e.g. CRM Data Warehouse")
    author_name = st.text_input("Author / Analyst", placeholder="e.g. Jane Smith")

    st.markdown('<div class="section-header" style="margin-top:1rem;">Documents to Generate</div>', unsafe_allow_html=True)

    col_c1, col_c2 = st.columns(2)
    with col_c1:
        gen_sttm    = st.checkbox("STTM", value=True)
        gen_dict    = st.checkbox("Data Dictionary", value=True)
        gen_lineage = st.checkbox("Data Lineage", value=True)
    with col_c2:
        gen_odd     = st.checkbox("ODD", value=True)
        gen_frd     = st.checkbox("FRD", value=True)
        gen_excel   = st.checkbox("Excel Bundle", value=True)

    st.markdown("<br>", unsafe_allow_html=True)
    generate = st.button("⬡  Generate Documentation", use_container_width=True)


# ---------------------------
# MAIN LOGIC
# ---------------------------

if generate:
    if not sql_input.strip():
        st.error("Please paste a SQL query to continue.")
    else:
        try:
            with st.spinner("Parsing SQL and building lineage…"):
                engine = LineageEngine(sql_input)
                lineage = engine.run()

            if not lineage:
                st.warning("No column lineage could be extracted. Check your SQL syntax.")
                st.stop()

            df = pd.DataFrame(lineage)

            # Metrics
            st.markdown("<br>", unsafe_allow_html=True)
            m1, m2, m3, m4 = st.columns(4)
            m1.metric("Source Tables", len(engine.get_source_tables()))
            m2.metric("Target Columns", len(engine.get_target_columns()))
            m3.metric("Lineage Rows", len(df))
            m4.metric("CTEs Found", len(engine.cte_map))

            st.markdown("<br>", unsafe_allow_html=True)

            # Tabs
            tabs = st.tabs(["📋 STTM", "📖 Data Dict", "🔗 Lineage", "📦 Downloads"])

            with tabs[0]:
                st.markdown('<div class="section-header">Source-to-Target Mapping</div>', unsafe_allow_html=True)
                display_cols = ["SOURCE_TABLE", "SOURCE_COLUMN", "TRANSFORMATION_TYPE", "TRANSFORMATION_LOGIC", "TARGET_COLUMN"]
                st.dataframe(df[display_cols], use_container_width=True, height=350)

            with tabs[1]:
                st.markdown('<div class="section-header">Data Dictionary Preview</div>', unsafe_allow_html=True)
                seen = set()
                dict_rows = []
                for _, row in df.iterrows():
                    col = row["TARGET_COLUMN"]
                    if col in seen:
                        continue
                    seen.add(col)
                    src = df[df["TARGET_COLUMN"] == col]["SOURCE_TABLE"].unique()
                    dict_rows.append({
                        "Column": col,
                        "Source Table(s)": ", ".join(s for s in src if s not in ("N/A","UNKNOWN")) or "Derived",
                        "Transform Type": row["TRANSFORMATION_TYPE"],
                        "Logic": row["TRANSFORMATION_LOGIC"],
                    })
                st.dataframe(pd.DataFrame(dict_rows), use_container_width=True, height=350)

            with tabs[2]:
                st.markdown('<div class="section-header">Lineage Summary</div>', unsafe_allow_html=True)
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown("**Source Tables**")
                    for t in engine.get_source_tables():
                        st.markdown(f"`{t}`")
                with c2:
                    st.markdown("**Transformation Types**")
                    type_counts = df["TRANSFORMATION_TYPE"].value_counts()
                    for k, v in type_counts.items():
                        st.markdown(f"`{k}` — {v} column(s)")

            with tabs[3]:
                st.markdown('<div class="section-header">Download Documents</div>', unsafe_allow_html=True)

                dl_col1, dl_col2, dl_col3 = st.columns(3)

                with dl_col1:
                    if gen_sttm:
                        st.download_button(
                            "⬇ STTM.docx",
                            data=generate_sttm_doc(df, project_name, author_name),
                            file_name="STTM.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    if gen_dict:
                        st.download_button(
                            "⬇ DataDictionary.docx",
                            data=generate_data_dict_doc(df, project_name, author_name),
                            file_name="DataDictionary.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                with dl_col2:
                    if gen_lineage:
                        st.download_button(
                            "⬇ DataLineage.docx",
                            data=generate_lineage_doc(df, engine, project_name, author_name),
                            file_name="DataLineage.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    if gen_odd:
                        st.download_button(
                            "⬇ ODD.docx",
                            data=generate_odd_doc(df, engine, project_name, author_name),
                            file_name="ODD.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                with dl_col3:
                    if gen_frd:
                        st.download_button(
                            "⬇ FRD.docx",
                            data=generate_frd_doc(df, engine, project_name, author_name),
                            file_name="FRD.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    if gen_excel:
                        st.download_button(
                            "⬇ SQLDocs_Bundle.xlsx",
                            data=generate_excel(df),
                            file_name="SQLDocs_Bundle.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )

                st.success(f"✓ Documentation generated — {datetime.now().strftime('%H:%M:%S')}")

        except Exception as e:
            st.error(f"Error: {e}")
